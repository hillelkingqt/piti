#!/usr/bin/env python3
"""
---------------------------------------------------------------------------------
 generate_office_docs.py  –  ⚙️ Advanced Office‑doc generator (v2.0.4, 2025‑05‑29)
---------------------------------------------------------------------------------
Changelog (v2.0.4)
• FIX: OMML equation insertion refined. Now directly appends parsed CT_OMath
  object to the paragraph's OXML element if the input OMML is a full <m:oMath> block.
  This is a more robust way to handle valid OMML from JSON.
• FIX: Strengthened default RTL and right-alignment for Hebrew documents.
  Explicitly sets overall document direction and ensures paragraph styles reflect RTL.
Changelog (v2.0.3)
• FIX: Corrected `doc.settings.element` usage.
• FIX: Improved `_insert_omml_equation` to handle wrapped/unwrapped OMML.
Changelog (v2.0.2)
• FIX: OMML equation insertion method changed.
• FIX: Ensured more consistent right alignment and RTL for Hebrew.
• FIX: SyntaxError in xpath calls.
(Previous v2.0 changes omitted for brevity)
---------------------------------------------------------------------------------
"""

from __future__ import annotations
import json
import sys
import logging
import traceback
import time
import errno
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple, Union, cast
import re
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt               # היה קיים
from docx.shared import Inches           # ← הוסף כאן
from docx.oxml.ns import qn   
# --- Optional imports with clear error handling ---
_MODULE_ERR_MSG = "Missing optional package '{name}'. Please install it by running: pip install {pip_name}"

try:
    import docx
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
    from docx.enum.style import WD_STYLE_TYPE
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn, nsdecls
    from docx.table import _Cell as DocxCellType 
    from docx.text.paragraph import Paragraph as DocxParagraphType 
    from docx.text.run import Run as DocxRunType 
except ModuleNotFoundError:
    docx = None # type: ignore
    Document = None # type: ignore
    Pt = Inches = RGBColor = None # type: ignore
    WD_ALIGN_PARAGRAPH = WD_BREAK = None # type: ignore
    WD_STYLE_TYPE = None # type: ignore
    OxmlElement = None # type: ignore
    qn = nsdecls = None # type: ignore
    DocxCellType = None # type: ignore
    DocxParagraphType = None # type: ignore
    DocxRunType = None # type: ignore


try:
    import openpyxl
    from openpyxl import Workbook
    from openpyxl.styles import Font, Alignment, Border, Side, PatternFill
    from openpyxl.utils import get_column_letter
    from openpyxl.worksheet.worksheet import Worksheet as OpenpyxlWorksheetType 
    from openpyxl.cell.cell import Cell as OpenpyxlCellType 
except ModuleNotFoundError:
    openpyxl = None # type: ignore
    Workbook = None # type: ignore
    Font = Alignment = Border = Side = PatternFill = None # type: ignore
    get_column_letter = None # type: ignore
    OpenpyxlWorksheetType = None # type: ignore
    OpenpyxlCellType = None # type: ignore


# ────────────────────────── Constants ─────────────────────────────────────────
DEFAULT_FONT_FAMILY = "Calibri"
DEFAULT_HEBREW_FONT_FAMILY = "David" 
FALLBACK_HEBREW_FONT_FAMILY = "Arial" 

DEFAULT_FONT_SIZE_PT = 11
DEFAULT_HEADING_FONT_SIZE_PT = {
    0: 24, 1: 16, 2: 14, 3: 13, 4: 12,
}

COLOR_BLACK = (0, 0, 0)
COLOR_WHITE = (255, 255, 255)
COLOR_BLUE = (0, 112, 192)

# ────────────────────────── Utils ─────────────────────────────────────────────
# _fatal, _ensure_module, _safe_save, _parse_color remain the same as v2.0.3
def _fatal(msg: str, exc: Optional[Exception] = None) -> None:
    logging.error(msg)
    if exc: logging.error(f"Exception details: {exc}")
    sys.exit(1)

def _ensure_module(module: Optional[Any], name: str, pip_name: str) -> None:
    if module is None: _fatal(_MODULE_ERR_MSG.format(name=name, pip_name=pip_name))

def _safe_save(save_callable, out_path: Path) -> Path:
    try:
        save_callable(out_path)
        logging.info(f"Document saved successfully to: '{out_path}'")
        return out_path
    except PermissionError as exc:
        if exc.errno not in (errno.EACCES, errno.EPERM): raise
        ts = int(time.time())
        alt_path = out_path.with_stem(f"{out_path.stem}_readonly_{ts}")
        logging.warning(f"PermissionError saving to '{out_path}'. Retrying with '{alt_path.name}'")
        try:
            save_callable(alt_path)
            logging.info(f"Document saved to alternative path: '{alt_path}'")
            return alt_path
        except Exception as final_exc: _fatal(f"Failed to save to '{alt_path.name}'.", final_exc)
    except Exception as e: _fatal(f"Error saving to '{out_path}'.", e)
    return out_path # Should be unreachable

def _parse_color(color_str: Optional[str]) -> Optional[RGBColor]:
    if not color_str or RGBColor is None: return None
    color_str = color_str.strip().upper()
    try:
        if color_str.startswith("RGB("):
            parts = color_str[4:-1].split(',')
            r, g, b = [int(p.strip()) for p in parts]
            return RGBColor(r, g, b)
        elif len(color_str) == 6:
            return RGBColor.from_string(color_str)
    except ValueError: pass
    logging.warning(f"Invalid color string: {color_str}. Using default.")
    return None

# ────────────────────────── Word: Style Helpers ───────────────────────────────
# _get_font_family, _apply_run_formatting, _apply_paragraph_formatting, _set_default_hebrew_styles
# remain largely the same as v2.0.3, with minor refinements for clarity or default handling.

def _get_font_family(el_data: Dict[str, Any], global_data: Dict[str, Any], is_hebrew: bool) -> str:
    family = el_data.get("font_family") or global_data.get("font_family")
    if family: return family
    return DEFAULT_HEBREW_FONT_FAMILY if is_hebrew else DEFAULT_FONT_FAMILY

def _apply_run_formatting(run: DocxRunType, el_data: Dict[str, Any], global_data: Dict[str, Any], is_hebrew: bool) -> None:
    if Pt is None: return 
    font = run.font
    font.name = _get_font_family(el_data, global_data, is_hebrew)
    if is_hebrew: font.cs_font_name = DEFAULT_HEBREW_FONT_FAMILY
    
    if (size_pt := el_data.get("font_size")): font.size = Pt(int(size_pt))
    elif not run.style or run.style.name == 'Default Paragraph Font': 
        font.size = Pt(global_data.get("font_size_pt", DEFAULT_FONT_SIZE_PT))

    if el_data.get("bold"): font.bold = True
    if el_data.get("italic"): font.italic = True
    if el_data.get("underline"): font.underline = True
    if el_data.get("strike"): font.strike = True
    if (color_val := _parse_color(el_data.get("font_color"))): font.color.rgb = color_val

def _apply_paragraph_formatting(
    paragraph: "DocxParagraphType",
    el_data: dict,
    global_data: dict,
    is_hebrew: bool,
) -> None:
    """
    Alignment + כיוון כתיבה (RTL/Bidi) כולל mirrorIndents לרשימות.
    """

    # ── Alignment ────────────────────────────────────────────────────────────
    align_str = el_data.get("alignment")
    if align_str is None:
        align_str = global_data.get("default_alignment")
    if align_str is None:
        align_str = "right" if is_hebrew else "left"

    if align_str == "center":
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align_str == "right":
        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    else:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # ── bidi / RTL  ─────────────────────────────────────────────────────────
    apply_rtl = el_data.get("rtl")
    if apply_rtl is None:
        apply_rtl = global_data.get("rtl_all", is_hebrew)

    p_pr = paragraph._p.get_or_add_pPr()        # <w:pPr>

    if apply_rtl:
        paragraph.paragraph_format.bidi = True
        p_pr.set(qn("w:bidi"), "1")             # Right-to-Left Text Direction
        p_pr.set(qn("w:mirrorIndents"), "1")    # הופך left indent → right indent
        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT  # מוודא יישור לימין
    else:
        paragraph.paragraph_format.bidi = False
        if p_pr.get(qn("w:bidi")):
            p_pr.attrib.pop(qn("w:bidi"))
        if p_pr.get(qn("w:mirrorIndents")):
            p_pr.attrib.pop(qn("w:mirrorIndents"))


def _set_default_hebrew_styles(doc: Document) -> None:
    """
    הופך את הסגנונות הראשיים של Word (Normal, Title, Heading 1-4) לכיוון ימין-לשמאל
    ומגדיר גופן עברי בגודל הגיוני.
    """
    DEFAULT_HEADING_FONT_SIZE_PT = {
        0: 24,  # Title
        1: 22,
        2: 18,
        3: 16,
        4: 14,
    }

    styles_to_modify = {
        "Normal": None,
        "Title": DEFAULT_HEADING_FONT_SIZE_PT[0],
    }
    for i in range(1, 5):
        styles_to_modify[f"Heading {i}"] = DEFAULT_HEADING_FONT_SIZE_PT[i]

    for style_name, font_size in styles_to_modify.items():
        style = doc.styles[style_name]
        pf = style.paragraph_format
        pf.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        pf.bidi = True
        if font_size is not None:
            style.font.size = Pt(font_size)
        style.font.name = "David"

def _insert_omml_equation(paragraph: DocxParagraphType, omml_input_str: str) -> None:
    """
    Revised OMML insertion.
    omml_input_str is expected to be a string representing a full <m:oMath>...</m:oMath> element.
    """
    if OxmlElement is None or qn is None or docx is None: return

    omml_input_str = omml_input_str.strip()

    # LaTeX → OMML  (כל מחרוזת שלא מתחילה ב-"<")
    if not omml_input_str.lstrip().startswith("<"):
        from latex2mathml.converter import convert as l2m
        import mathml2omml
        omml_input_str = mathml2omml.convert(l2m(omml_input_str))


    # עיטוף אוטומטי אם חסר <m:oMath>
    if "<m:oMath" not in omml_input_str:
        omml_input_str = (
            '<m:oMath xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math">'
            + omml_input_str +
            "</m:oMath>"
        )

    # הזרקת namespace אם עדיין חסר
    if omml_input_str.startswith("<m:oMath") and "xmlns:m=" not in omml_input_str:
        omml_input_str = omml_input_str.replace(
            "<m:oMath",
            '<m:oMath xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math"',
            1,
        )


    # ► הזרקת namespace אם חסר
    if omml_input_str.startswith("<m:oMath") and "xmlns:m=" not in omml_input_str:
        omml_input_str = omml_input_str.replace(
            "<m:oMath",
            '<m:oMath xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math"',
            1
        )
    if not omml_input_str.startswith("<m:oMath"):
        logging.warning(f"OMML input does not start with <m:oMath>. Attempting to use as is, but it might fail. Input: {omml_input_str[:100]}")
        # Fallback to old method if it's just a fragment, though this might not be ideal.
        # For now, we rely on the input being a full <m:oMath> element.
        # If it's a fragment, the user of this script (Node.js) should wrap it.

    try:
        # Parse the string to an OXML element
        omml_el = docx.oxml.parse_xml(omml_input_str)
        
        # The <m:oMath> element should be directly appended to the paragraph's <w:p> element.
        # It is typically wrapped in a <w:r> for display, but appending to _p seems to work
        # for block-level math. For inline math, it would be inside a run.
        # For block display, often it's the only content of a paragraph.
        # We'll add it to the paragraph's OXML structure.
        paragraph._p.append(omml_el)
        logging.info("Successfully appended OMML object to paragraph.")

    except Exception as e: # Catch lxml.etree.XMLSyntaxError or other parsing issues
        logging.error(f"Failed to parse or append OMML: {e}. Input (start): {omml_input_str[:150]}...")
        # Add a visible error message in the document
        err_run = paragraph.add_run(f"[OMML Parse Error: Could not process equation '{omml_input_str[:50]}...']")
        err_run.italic = True
        err_run.font.color.rgb = RGBColor(255,0,0) if RGBColor else None


# ────────────────────────── Word: Element Builders ───────────────────────────
# _add_paragraph_element, _add_heading_element, etc. remain the same as v2.0.3
def _add_paragraph_element(doc: Document, el_data: Dict[str, Any],
                           global_data: Dict[str, Any], is_hebrew: bool) -> None:
    """
    מוסיף פסקה; אם יש בה ביטויי $…$ או \( … \) – ממיר כל אחד ל-OLMM בתור נוסחה.
    """
    text = el_data.get("text", "")
    p = doc.add_paragraph()
    _apply_paragraph_formatting(p, el_data, global_data, is_hebrew)

    # אין LaTeX? פשוט הוסף Run רגיל
    if "$" not in text and "\\(" not in text:
        run = p.add_run(text)
        _apply_run_formatting(run, el_data, global_data, is_hebrew)
        return

    # יש LaTeX → מפרקים לטוקנים: טקסט רגיל | LaTeX תחום
    tokens = re.split(r'(\\[A-Za-z]+[^\\$]*|\\\([^\)]+\\\)|\$[^$]+\$)', text)
    for tk in tokens:
        if not tk:
            continue
        if tk.startswith("$") and tk.endswith("$"):
            latex_expr = tk[1:-1]              # בלי $ … $
            _insert_omml_equation(p, latex_expr)
        elif tk.startswith("\\(") and tk.endswith("\\)"):
            latex_expr = tk[2:-2]              # בלי \( … \)
            _insert_omml_equation(p, latex_expr)
        else:
            run = p.add_run(tk)
            _apply_run_formatting(run, el_data, global_data, is_hebrew)

def _add_heading_element(doc: Document, el_data: Dict[str, Any], global_data: Dict[str, Any], is_hebrew: bool) -> None:
    if WD_ALIGN_PARAGRAPH is None: return
    text = el_data.get("text", "")
    level = int(el_data.get("level", 1))
    h = doc.add_heading(level=level) 
    # Headings are also paragraphs, so clear existing runs from add_heading if we want to set text manually.
    for r in h.runs: r.clear() # Clear default text run if any
    _apply_paragraph_formatting(h, el_data, global_data, is_hebrew) 
    run = h.add_run(text) 
    _apply_run_formatting(run, el_data, global_data, is_hebrew)

def _add_list_item_element(
    doc: Document,
    el_data: Dict[str, Any],
    global_data: Dict[str, Any],
    is_hebrew: bool,
) -> None:
    """
    פריט רשימה RTL אמיתי – הנקודה/מספר בצד ימין, טקסט תלוי.
    """
    text = el_data.get("text", "")
    style_type = el_data.get("style", "bullet").lower()
    style_name = "List Number" if style_type == "number" else "List Bullet"

    # ─ 1. פסקה בסגנון הרשימה
    try:
        p = doc.add_paragraph(style=style_name)
    except KeyError:
        p = doc.add_paragraph()             # fallback
        p.add_run(("• " if style_type == "bullet" else "1. ") + text)
    else:
        p.add_run(text)

    # ─ 2. RTL, יישור, mirror
    _apply_paragraph_formatting(p, el_data, global_data, is_hebrew)
    p_pr = p._p.get_or_add_pPr()
    p_pr.set(qn("w:bidi"), "1")             # כיוון ימין-לשמאל
    p_pr.set(qn("w:mirrorIndents"), "1")    # הופך left→right ברשימה

    # ─ 3. תלייה: LEFT + HANGING (לא RIGHT!)
    ind = p_pr.get_or_add_ind()
    ind.set(qn("w:left"),  "360")           # 360 twips ≈ 0.25″
    ind.set(qn("w:hanging"), "360")         # שורה ראשונה תלויה

    # ─ 4. עיצוב Run-ים
    for run in p.runs:
        _apply_run_formatting(run, el_data, global_data, is_hebrew)


def _add_image_element(doc: Document, el_data: Dict[str, Any]) -> None:
    if Inches is None: return
    img_path_str = el_data.get("path")
    if not img_path_str: logging.warning("Image: 'path' not provided."); return
    img_path = Path(img_path_str)
    if not img_path.is_file(): logging.warning(f"Image: File not found '{img_path}'."); return
    try:
        w_in = float(el_data["width_inches"]) if "width_inches" in el_data else None
        h_in = float(el_data["height_inches"]) if "height_inches" in el_data else None
        p = doc.add_paragraph()
        _apply_paragraph_formatting(p, el_data, {}, el_data.get("rtl", False))
        run = p.add_run() 
        if w_in and h_in: run.add_picture(str(img_path), width=Inches(w_in), height=Inches(h_in))
        elif w_in: run.add_picture(str(img_path), width=Inches(w_in))
        elif h_in: run.add_picture(str(img_path), height=Inches(h_in))
        else: run.add_picture(str(img_path))
        logging.info(f"Image '{img_path.name}' added.")
    except Exception as exc: logging.warning(f"Image '{img_path.name}' skipped: {exc}")

def _add_table_element(doc: Document, el_data: Dict[str, Any], global_data: Dict[str, Any], is_hebrew: bool) -> None:
    if OxmlElement is None or qn is None: return
    table_data = el_data.get("data")
    if not (table_data and isinstance(table_data, list) and all(isinstance(r, list) for r in table_data)):
        logging.warning("Table: 'data' missing or not list of lists."); return
    rows, cols = len(table_data), len(table_data[0]) if rows > 0 else 0
    if not (rows and cols): logging.warning("Table: data is empty."); return
    table = doc.add_table(rows=rows, cols=cols)
    table.style = el_data.get("style", "Table Grid") 
    table_is_rtl = el_data.get("rtl_table", is_hebrew)
    if table_is_rtl:
        tblPr = table._tbl.get_or_add_tblPr()
        if not tblPr.find(qn('w:bidiVisual')): tblPr.append(OxmlElement('w:bidiVisual'))
    for i, row_dt in enumerate(table_data):
        for j, cell_txt in enumerate(row_dt):
            if j < cols: 
                cell = table.cell(i, j)
                p_in_cell = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
                p_in_cell.clear()
                run = p_in_cell.add_run(str(cell_txt))
                cell_style = el_data.get("cell_style", {})
                _apply_run_formatting(run, cell_style, global_data, table_is_rtl)
                _apply_paragraph_formatting(p_in_cell, cell_style, global_data, table_is_rtl)
    col_w = el_data.get("col_widths")
    if col_w and isinstance(col_w, list) and len(col_w) == cols:
        for j, w_val in enumerate(col_w):
            try: table.columns[j].width = Inches(float(w_val))
            except ValueError: logging.warning(f"Invalid column width: {w_val}")
    logging.info(f"Table {rows}x{cols} added.")

def _add_page_break_element(doc: Document) -> None:
    if WD_BREAK is None: return
    doc.add_page_break(); logging.info("Page break added.")

def _add_hyperlink_element(doc: Document, p: DocxParagraphType, el_data: Dict[str, Any], global_data: Dict[str, Any], is_hebrew: bool) -> None:
    if OxmlElement is None or qn is None or RGBColor is None or docx is None: return
    url, text = el_data.get("url"), el_data.get("text")
    if not url: logging.warning("Hyperlink: 'url' not provided."); p.add_run("[Invalid Hyperlink]"); return
    text = text or url
    r_id = p.part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)
    hyperlink, sub_run_el = OxmlElement('w:hyperlink'), OxmlElement('w:r')
    hyperlink.set(qn('r:id'), r_id)
    rPr_el = OxmlElement('w:rPr')
    h_style_name = "Hyperlink"
    if el_data.get("style", h_style_name) == h_style_name and h_style_name in doc.styles:
        style_el = OxmlElement('w:rStyle'); style_el.set(qn('w:val'), h_style_name); rPr_el.append(style_el)
    else:
        dummy_p = doc.add_paragraph(); temp_run = dummy_p.add_run()
        _apply_run_formatting(temp_run, el_data, global_data, is_hebrew)
        if temp_run.font.color.rgb is None and COLOR_BLUE: temp_run.font.color.rgb = RGBColor(*COLOR_BLUE)
        temp_run.font.underline = True
        if temp_run._r.rPr: rPr_el = temp_run._r.rPr.copy()
        doc.element.body.remove(dummy_p._p)
    sub_run_el.append(rPr_el)
    t_el = OxmlElement('w:t'); t_el.text = text; sub_run_el.append(t_el)
    hyperlink.append(sub_run_el); p._p.append(hyperlink)
    logging.info(f"Hyperlink '{text}' to '{url}' added.")

def _add_header_footer_element(doc: Document, el_data: Dict[str, Any], global_data: Dict[str, Any], is_hebrew: bool, section_idx: int = 0, is_header: bool = True) -> None:
    if WD_ALIGN_PARAGRAPH is None or OxmlElement is None or qn is None: return
    section = doc.sections[section_idx]
    container = section.header if is_header else section.footer
    p = container.paragraphs[0] if container.paragraphs else container.add_paragraph()
    p.clear()
    _apply_paragraph_formatting(p, el_data, global_data, is_hebrew)
    text, page_num = el_data.get("text", ""), el_data.get("page_number", False)
    if page_num:
        fldSimple = OxmlElement('w:fldSimple'); fldSimple.set(qn('w:instr'), r' PAGE \* MERGEFORMAT ')
        run_pg = OxmlElement('w:r'); run_pg.append(fldSimple); p._p.append(run_pg)
        if text: p.add_run(" ")
    if text:
        run_txt = p.add_run(text)
        _apply_run_formatting(run_txt, el_data.get("style",{}), global_data, is_hebrew)
    logging.info(f"{'Header' if is_header else 'Footer'} content added.")

# ────────────────────────── Word: Document Builder ────────────────────────────

def _build_word(data: Dict[str, Any]) -> str:
    _ensure_module(docx, "python-docx", "python-docx")
    if Document is None or OxmlElement is None or qn is None: _fatal("Word deps not loaded."); return ""

    doc = Document()
    doc_lang = data.get("doc_lang", "en").lower()
    is_hebrew = (doc_lang == "he")
    global_data = data.get("global_settings", {})
    global_data["font_size_pt"] = global_data.get("font_size_pt", DEFAULT_FONT_SIZE_PT)

    section = doc.sections[0] # Page setup, margins, etc. (same as v2.0.3)
    if "margins" in global_data:
        mg = global_data["margins"]
        if "top_inches" in mg: section.top_margin = Inches(float(mg["top_inches"]))
        if "bottom_inches" in mg: section.bottom_margin = Inches(float(mg["bottom_inches"]))
        if "left_inches" in mg: section.left_margin = Inches(float(mg["left_inches"]))
        if "right_inches" in mg: section.right_margin = Inches(float(mg["right_inches"]))
    if global_data.get("orientation", "portrait").lower() == "landscape":
        section.orientation = docx.enum.section.WD_ORIENTATION.LANDSCAPE
        section.page_width, section.page_height = section.page_height, section.page_width

    if is_hebrew:
        # Set overall document to RTL
        settings = doc.settings.element # This is <w:settings>
        doc_grid = settings.find(qn('w:docGrid'))
        if doc_grid is not None : # Try to insert bidi before docGrid for better compatibility
             idx = settings.index(doc_grid)
        else: # append to settings
             idx = len(settings)

        # Ensure <w:bidi/> exists directly under <w:settings> for document-level RTL view
        bidi_setting_el = settings.find(qn('w:bidi'))
        if bidi_setting_el is None:
            bidi_setting_el = OxmlElement('w:bidi')
            settings.insert(idx,bidi_setting_el) # Insert or append
        # bidi_setting_el.set(qn('w:val'), '1') # Often not needed if element present
        
        # For default text direction within paragraphs (docDefaults)
        doc_defaults = settings.find(qn('w:docDefaults'))
        if doc_defaults is None:
            doc_defaults = OxmlElement('w:docDefaults')
            settings.insert(idx, doc_defaults) # Insert or append
        
        rpr_defaults = doc_defaults.find(qn('w:rPrDefault'))
        if rpr_defaults is None:
            rpr_defaults = OxmlElement('w:rPrDefault')
            doc_defaults.append(rpr_defaults)
        
        rpr = rpr_defaults.find(qn('w:rPr'))
        if rpr is None:
            rpr = OxmlElement('w:rPr')
            rpr_defaults.append(rpr)

        if rpr.find(qn('w:rtl')) is None: rpr.append(OxmlElement('w:rtl'))
        
        lang_el = rpr.find(qn('w:lang'))
        if lang_el is None: lang_el = OxmlElement('w:lang'); rpr.append(lang_el)
        lang_el.set(qn('w:val'), 'he-IL'); lang_el.set(qn('w:eastAsia'), 'en-US'); lang_el.set(qn('w:bidi'), 'he-IL')
        
        if section._sectPr.find(qn('w:rtlGutter')) is None:
            sectPr_el = section._sectPr
            rtlgutter_el = OxmlElement('w:rtlGutter'); rtlgutter_el.set(qn('w:val'), 'true')
            sectPr_el.append(rtlgutter_el)
            
        _set_default_hebrew_styles(doc)

    if (title_text := data.get("title")):
        title_data = {"text": title_text, "level": 0, "alignment": "center" if not is_hebrew else "right"}
        _add_heading_element(doc, title_data, global_data, is_hebrew)

    # Headers/Footers (same as v2.0.3)
    if (hd := data.get("header")): _add_header_footer_element(doc, hd, global_data, is_hebrew, is_header=True)
    if (ft := data.get("footer")): _add_header_footer_element(doc, ft, global_data, is_hebrew, is_header=False)
    elif global_data.get("page_numbers", is_hebrew): # Default page numbers true for Hebrew
        df_align = "right" if is_hebrew else "center"
        _add_header_footer_element(doc, {"page_number": True, "alignment": df_align}, global_data, is_hebrew, is_header=False)

    for el_data in data.get("content", []):
        el_type = el_data.get("type", "").lower()
        try:
            if el_type == "paragraph": _add_paragraph_element(doc, el_data, global_data, is_hebrew)
            elif el_type == "heading": _add_heading_element(doc, el_data, global_data, is_hebrew)
            elif el_type == "omml_math" or el_type == "equation": 
                p = doc.add_paragraph() 
                eq_align = el_data.get("alignment", "center" if not is_hebrew else "right") # Default right for HE math
                _apply_paragraph_formatting(p, {**el_data, "alignment": eq_align}, global_data, is_hebrew) 
                if (omml := el_data.get("omml")): _insert_omml_equation(p, omml)
                else: logging.warning("Equation: 'omml' missing."); p.add_run("[Empty Equation]").italic = True
            elif el_type == "list_item": _add_list_item_element(doc, el_data, global_data, is_hebrew)
            elif el_type == "image": _add_image_element(doc, el_data)
            elif el_type == "table": _add_table_element(doc, el_data, global_data, is_hebrew)
            elif el_type == "page_break": _add_page_break_element(doc)
            elif el_type == "hyperlink":
                p = doc.add_paragraph() 
                _apply_paragraph_formatting(p, el_data, global_data, is_hebrew)
                _add_hyperlink_element(doc, p, el_data, global_data, is_hebrew)
            else: logging.warning(f"Unknown element type: '{el_type}'")
        except Exception as e:
            logging.error(f"Failed processing element '{el_type}': {e}. Data: {str(el_data)[:100]}"); traceback.print_exc()

    out_path = Path(data["output_path"])
    out_path.parent.mkdir(parents=True, exist_ok=True)
    return str(_safe_save(doc.save, out_path))

# ────────────────────────── Excel: Helpers & Builder ──────────────────────────
# (Excel part remains the same as v2.0.3)
def _apply_excel_cell_style(cell: OpenpyxlCellType, style_data: Dict[str, Any], global_data: Dict[str, Any], is_hebrew: bool) -> None:
    if Font is None or Alignment is None or Border is None or Side is None or PatternFill is None or RGBColor is None: return
    font_family = _get_font_family(style_data, global_data, is_hebrew)
    font_size = style_data.get("font_size", global_data.get("font_size_pt", 11))
    is_bold, is_italic = style_data.get("bold", False), style_data.get("italic", False)
    font_color_hex = None
    if (fcs := style_data.get("font_color")) and (frgb := _parse_color(fcs)): font_color_hex = f"FF{frgb.r:02X}{frgb.g:02X}{frgb.b:02X}"
    cell.font = Font(name=font_family, size=font_size, bold=is_bold, italic=is_italic, color=font_color_hex)
    h_align = style_data.get("horizontal_alignment", "right" if is_hebrew else "general")
    v_align = style_data.get("vertical_alignment", "bottom")
    cell.alignment = Alignment(horizontal=h_align, vertical=v_align, wrap_text=style_data.get("wrap_text", False))
    if "border" in style_data:
        bs, bc_str = style_data["border"].get("style", "thin"), style_data["border"].get("color", "000000")
        bc_hex = f"FF{_parse_color(bc_str).r:02X}{_parse_color(bc_str).g:02X}{_parse_color(bc_str).b:02X}" if _parse_color(bc_str) else "FF000000"
        side = Side(style=bs, color=bc_hex)
        cell.border = Border(left=side, right=side, top=side, bottom=side)
    if (fill_cs := style_data.get("fill_color")) and (fill_rgb := _parse_color(fill_cs)):
        fill_hex = f"FF{fill_rgb.r:02X}{fill_rgb.g:02X}{fill_rgb.b:02X}"
        cell.fill = PatternFill(start_color=fill_hex, end_color=fill_hex, fill_type="solid")
    if (nf := style_data.get("number_format")): cell.number_format = nf

def _build_excel(data: Dict[str, Any]) -> str:
    _ensure_module(openpyxl, "openpyxl", "openpyxl")
    if Workbook is None or Font is None or Alignment is None or get_column_letter is None: _fatal("Excel deps not loaded."); return ""
    wb = Workbook()                                   # יצירת חוברת
    ws = cast(OpenpyxlWorksheetType, wb.active)       # הגיליון הפעיל
    ws.title = data.get("sheet_name", "Sheet1")
    is_hebrew = data.get("doc_lang", "en").lower() == "he"
    global_data = data.get("global_settings", {})
    if data.get("sheet_rtl", is_hebrew): ws.sheet_view.rightToLeft = True
    row_num = 1
    if (title := data.get("title")):
        cols_merge = len(data.get("excel_data", [[]])[0]) if data.get("excel_data") else 1
        if cols_merge > 0: ws.merge_cells(start_row=row_num, start_column=1, end_row=row_num, end_column=cols_merge)
        title_cell = ws.cell(row=row_num, column=1, value=title)
        title_style = data.get("title_style", {"bold": True, "font_size": 14, "horizontal_alignment": "center"})
        _apply_excel_cell_style(cast(OpenpyxlCellType, title_cell), title_style, global_data, is_hebrew)
        row_num += 2
    excel_data = data.get("excel_data")
    if not (excel_data and isinstance(excel_data, list)): logging.warning("Excel: 'excel_data' missing/invalid.")
    else:
        hdr_style = data.get("header_style", {"bold": True, "horizontal_alignment": "center" if not is_hebrew else "right"})
        for r_idx, row_val in enumerate(excel_data):
            if not isinstance(row_val, list): logging.warning(f"Row {r_idx+1} not list, skipping."); continue
            max_c = 0
            for c_idx, cell_val in enumerate(row_val):
                cell = ws.cell(row=row_num + r_idx, column=c_idx + 1, value=cell_val)
                _apply_excel_cell_style(cast(OpenpyxlCellType, cell), hdr_style if r_idx == 0 else {}, global_data, is_hebrew)
                max_c = c_idx
            if data.get("auto_adjust_column_widths", True):
                for c_adj in range(max_c + 1): ws.column_dimensions[get_column_letter(c_adj + 1)].auto_size = True
    if (col_w := data.get("column_widths")) and isinstance(col_w, list):
        for i, w in enumerate(col_w):
            try: ws.column_dimensions[get_column_letter(i + 1)].width = float(w)
            except ValueError: logging.warning(f"Invalid Excel col width: {w}")
    out_path = Path(data["output_path"])
    out_path.parent.mkdir(parents=True, exist_ok=True)
    return str(_safe_save(wb.save, out_path))

# ────────────────────────── Driver ──────────────────────────────────────────
# (Driver section remains the same as v2.0.3)
def _derive_out_path(data: Dict[str, Any], inp_path: Path) -> Path:
    if (p_str := data.get("output_path")): return Path(p_str)
    base_name = data.get("output_filename") or f"generated_document_{int(time.time())}"
    doc_type = data.get("doc_type", "word").lower()
    ext = ".docx" if doc_type == "word" else ".xlsx" if doc_type == "excel" else ".unknown"
    return inp_path.parent / (base_name + ext)

def main():
    logging.basicConfig(level=logging.INFO, format="[%(asctime)s - %(levelname)s] %(message)s", datefmt="%Y-%m-%d %H:%M:%S")
    if not (2 <= len(sys.argv) <= 3): print("Usage: python script.py <input.json> [output.docx/xlsx]"); sys.exit(1)
    inp_json_path = Path(sys.argv[1])
    if not inp_json_path.is_file(): _fatal(f"Input JSON not found: {inp_json_path}")
    try: data: Dict[str, Any] = json.loads(inp_json_path.read_text("utf-8"))
    except Exception as exc: _fatal(f"Error reading/parsing JSON '{inp_json_path.name}'", exc)
    data["output_path"] = str(Path(sys.argv[2] if len(sys.argv) == 3 else _derive_out_path(data, inp_json_path)).resolve())
    doc_type = data.get("doc_type", "word").lower()
    out_file_str = ""
    try:
        if doc_type == "word": out_file_str = _build_word(data)
        elif doc_type == "excel": out_file_str = _build_excel(data)
        else: _fatal(f"Unsupported doc_type: '{doc_type}'.")
    except Exception as exc: _fatal(f"Generation failed for '{doc_type}'", exc)
    if not out_file_str: _fatal("Generation resulted in empty output path.")
    sys.stdout.buffer.write((out_file_str + "\n").encode("utf-8", "replace")); sys.stdout.flush()
    logging.info(f"Successfully generated: {out_file_str}")

if __name__ == "__main__":
    main()