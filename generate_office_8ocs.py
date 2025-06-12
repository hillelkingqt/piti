
Gemini
What can Gemini do
in Google Drive
Summarize a topic
based on files in my Drive
Summarize a folder
in my Drive
Gemini for Workspace can make mistakes, including about people, so double-check it. Learn more
#!/usr/bin/env python3
"""
---------------------------------------------------------------------------------
 generate_office_docs.py  –  ⚙️ Office‑doc generator (v1.6, 2025‑05‑28)
---------------------------------------------------------------------------------
Changelog (v1.6)
• FIX: previous version was truncated → SyntaxError.  File now completes the
  final write/flush statement properly.
• Other behaviour unchanged:
    – Auto‑RTL when `doc_lang == "he"` (paragraphs, headings, lists, default
      right alignment).
    – Safe filename fallback on `PermissionError`.
    – UTF‑8‑safe STDOUT (writes bytes via `stdout.buffer`).
---------------------------------------------------------------------------------
"""

from __future__ import annotations
import json, sys, logging, traceback, time, errno
from pathlib import Path
from typing import Any

# ──────────────────────────  Utils  ────────────────────────────────────────────

def _fatal(msg: str, exc: Exception | None = None):
    logging.error(msg)
    if exc:
        traceback.print_exception(exc)
    sys.exit(1)


def _need(module, pip_name: str):
    if module is None:
        _fatal(f"Missing package '{pip_name}'. Run: pip install {pip_name}")


def _safe_save(save_callable, out_path: Path):
    """Try to save; on PermissionError retry with time‑stamped filename."""
    try:
        save_callable(out_path)
        return out_path
    except PermissionError as exc:
        if exc.errno not in (errno.EACCES, errno.EPERM):
            raise
        ts = int(time.time())
        alt = out_path.with_stem(out_path.stem + f"_{ts}")
        logging.warning(f"PermissionError saving '{out_path}'. Retrying with '{alt.name}' ...")
        save_callable(alt)
        return alt

# ──────────────────────────  Optional imports  ────────────────────────────────
try:
    import docx
    from docx.oxml import parse_xml
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ModuleNotFoundError:
    docx = None

try:
    import openpyxl
    from openpyxl.styles import Alignment, Font
except ModuleNotFoundError:
    openpyxl = None

# ──────────────────────────  Word helpers  ─────────────────────────────────────

def _style_run(run, element: dict[str, Any], global_data: dict[str, Any]):
    if element.get("bold"):
        run.bold = True
    if element.get("italic"):
        run.italic = True
    if (fs := element.get("font_size")):
        run.font.size = Pt(fs)
    if (fam := element.get("font_family") or global_data.get("font_family")):
        run.font.name = fam


def _insert_omml(paragraph, inner_xml: str):
    math_xml = (
        '<m:oMath xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math">'
        f"{inner_xml}" '</m:oMath>'
    )
    paragraph._p.append(parse_xml(math_xml))


def _apply_rtl(paragraph, el: dict[str, Any], doc_lang: str):
    """החלת כיוון ימין-לשמאל + יישור-ימין על פסקה נתונה (כולל כותרות, רשימות וכו')."""
    rtl = el.get("rtl") or doc_lang == "he"
    if not rtl:
        return

    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn

    # כיוון כתיבה RTL
    try:
        paragraph.paragraph_format.right_to_left = True
    except AttributeError:                      # fallback אם המאפיין חסר
        paragraph._p.get_or_add_pPr().set(qn("w:bidi"), "1")

    # תמיד מיושר לימין
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT


# ──────────────────────────  Word builder  ────────────────────────────────────

def _build_word(data: dict[str, Any]) -> str:
    _need(docx, "python-docx")
    doc = docx.Document()
    lang = data.get("doc_lang", "")

    # ――――――――― RTL + יישור-ימין ברירת-מחדל ―――――――――
    if lang == "he":
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn

        # כיוון כל ה-sections ל-RTL
        for s in doc.sections:
            try:
                s.right_to_left = True
            except AttributeError:
                pass

        # סגנון Normal – מיושר לימין וב-RTL
        normal = doc.styles["Normal"]
        normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        try:
            normal.paragraph_format.right_to_left = True
        except AttributeError:          # fallback לגרסאות ישנות של python-docx
            normal.element.get_or_add_pPr().set(qn("w:bidi"), "1")

    # ――――――――― תוכן המסמך ―――――――――
    if (title := data.get("title")):
        h = doc.add_heading(title, level=0)
        _apply_rtl(h, {}, lang)

    for el in data.get("content", []):
        et = el.get("type")
        if et == "paragraph":
            p = doc.add_paragraph(el.get("text", ""))
            _apply_rtl(p, el, lang)

        elif et == "heading":
            level = el.get("level", 1)
            h = doc.add_heading(el.get("text", ""), level=level)
            _apply_rtl(h, el, lang)

        elif et == "equation":
            p = doc.add_paragraph()
            _insert_omml(p, el.get("omml", ""))
            _apply_rtl(p, el, lang)

        elif et == "list_item":
            style = "List Bullet" if el.get("style", "bullet") == "bullet" else "List Number"
            p = doc.add_paragraph(el.get("text", ""), style=style)
            _apply_rtl(p, el, lang)

        elif et == "image":
            try:
                w = el.get("width_inches")
                doc.add_picture(el["path"], width=Inches(w) if w else None)
            except Exception as exc:
                logging.warning(f"Image skipped ({el.get('path')}): {exc}")

        else:
            logging.warning(f"Unknown element type skipped: {et}")

    out_path = Path(data["output_path"])
    out_path.parent.mkdir(parents=True, exist_ok=True)
    final_path = _safe_save(doc.save, out_path)
    return str(final_path)

# ──────────────────────────  Excel builder  ───────────────────────────────────

def _build_excel(data: dict[str, Any]) -> str:
    _need(openpyxl, "openpyxl")
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = data.get("sheet_name", "Sheet1")
    ws.sheet_view.rightToLeft = bool(data.get("sheet_rtl")) or data.get("doc_lang") == "he"

    row = 1
    if (title := data.get("title")):
        ws.merge_cells(start_row=1, start_column=1, end_row=1, end_column=len(data["excel_data"][0]))
        cell = ws.cell(row=1, column=1, value=title)
        cell.font = Font(size=14, bold=True)
        cell.alignment = Alignment(horizontal="center")
        row += 1

    for r in data["excel_data"]:
        for col, val in enumerate(r, 1):
            ws.cell(row=row, column=col, value=val)
        row += 1

    out_path = Path(data["output_path"])
    out_path.parent.mkdir(parents=True, exist_ok=True)
    final_path = _safe_save(wb.save, out_path)
    return str(final_path)

# ──────────────────────────  Driver  ──────────────────────────────────────────

def _derive_out_path(data: dict[str, Any], inp: Path) -> str:
    if p := data.get("output_path"):
        return p
    ts = int(time.time())
    stem = data.get("output_filename", f"generated_document_{ts}")
    ext = ".docx" if data.get("doc_type") == "word" else ".xlsx"
    return str(inp.with_name(stem + ext))


def main():
    logging.basicConfig(level=logging.INFO, format="[%(levelname)s] %(message)s")

    if len(sys.argv) not in (2, 3):
        _fatal("Usage: python generate_office_docs.py <input_json> [output_path]")

    inp_json = Path(sys.argv[1])
    if not inp_json.exists():
        _fatal(f"Input JSON not found: {inp_json}")

    try:
        data = json.loads(inp_json.read_text("utf-8"))
    except json.JSONDecodeError as exc:
        _fatal("Invalid JSON syntax", exc)

    data["output_path"] = sys.argv[2] if len(sys.argv) == 3 else _derive_out_path(data, inp_json)

    try:
        if data.get("doc_type") == "word":
            out_path = _build_word(data)
        elif data.get("doc_type") == "excel":
            out_path = _build_excel(data)
        else:
            _fatal(f"Unsupported doc_type: {data.get('doc_type')}")
    except Exception as exc:
        _fatal("Generation failed", exc)

    # Write UTF‑8 bytes so parent process need not decode console encoding
    sys.stdout.buffer.write((out_path + "\n").encode("utf-8", "replace"))
    sys.stdout.flush()

if __name__ == "__main__":
    main()